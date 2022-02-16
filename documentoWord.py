from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from datetime import datetime


document = Document()


dia= (datetime.today().strftime('%d-%m-%Y'))
styles = document.styles

# stilo parágrafo
p = styles.add_style('paragrafo', WD_STYLE_TYPE.PARAGRAPH)
p.font.name = 'Arial'
p.font.size = Pt(11)

# stilo heading2
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = styles["Heading 2"]
h2.font.name = "Arial"
h2.font.size = Pt(11)
h2.font.color.rgb = RGBColor(0, 0, 0)

a = styles.add_style("assinatura", WD_STYLE_TYPE.PARAGRAPH)
a.font.name = "Arial"
a.font.size = Pt(10)
a.font.color.rgb = RGBColor(0, 0, 0)
a.font.bold = True

r = styles.add_style("rodape", WD_STYLE_TYPE.PARAGRAPH)
r.font.name = "Times New Roman"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0, 0, 0)

valor= str(300)

document.add_paragraph('RECIBO                                                 VALOR: R$'+valor+',00', style="H2").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.add_paragraph('\n\nRecebi de '++', CPF' xxxxxxx', o valor de ('xxxxxx' reais), referente a 'x' sessões de Terapia Ocupacional a 'xxxxxxxxxx', dias ('xxxxxxx') de agosto de 2021.', style="paragrafo")

document.add_paragraph('\nSão josé dos Campos,(data).', style="paragrafo").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

document.add_paragraph('\n___________________________\nLucia Koyama de Jesus Silva\nTerapeuta Ocupacional\nCREFITO: 6351-TO\nCPF: 100.9292.88-97', style="assinatura")

document.add_paragraph('\n\nAlameda José Alves de Siqueira Filho, 52. Vila Betânia. CEP: 12245-492 São josé dos Campos-SP\ncel.(12)98812-9174 / e-mail. luciakoyama@outlook.com ', style="rodape").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.save('NewDocx.docx')