from Valida_cpf import ValidaCpf
from teste import Ja_tem
from teste import adicionar_criança
from PySimpleGUI import PySimpleGUI as sg
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from datetime import datetime
import time



#Layout

def janela_menu():
    sg.theme('Reddit')
    layout = [
        [sg.Text('O que deseja fazer?')],
        [sg.Button('1- Fazer recibos')],
        [sg.Button('2- Adicionar criança')],
        [sg.Button('3- Mostrar lista de crianças')],
        [sg.Button('4- Remover criança')]
    ]
    return sg.Window('Menu', layout=layout, finalize=True)

def janela_fazer():
    sg.theme('Reddit')
    Fazer = [
        [sg.Text('Para fazer o recibo insira a quantidade total de dias que a criança compareceu a terapia')],
        [sg.Text("Insira também os dias que ela compareceu da seguinte forma '01, 02, 03, 04...'")],
        [sg.Button('Voltar'), sg.Button('OK', key='do')]
    ]
    return sg.Window('Fazer Recibo', layout=Fazer, finalize=True)

def janela_continuarfazer():
    sg.theme('Reddit')
    Continuar = [
        [sg.Text('Quantos dias a criança veio a terepia?'), sg.Input(key='dias')],
        [sg.Text('Quais dias essa criança veio a terapia?'), sg.Input(key='quais')],
        [sg.Output(size=(30,20))],
        [sg.Button('Voltar'), sg.Button('Proximo', key='proximo'), sg.Button('Finalizar', key='finalizar')]
    ]
    return sg.Window('Fazer Recibo', layout=Continuar, finalize=True)

def janela_adicionar():
    sg.theme('Reddit')
    Adicionar = [
        [sg.Text('Nome do responsável:'), sg.Input(key='nomepai')],
        [sg.Text('CPF do responsável:'), sg.Input(key='cpfpai')],
        [sg.Text('Nome da criaça:'), sg.Input(key='nomecrianca')],
        #[sg.Text('CPF da criança:'), sg.Input(key='cpfcrianca')],
        [sg.Text('Valor'), sg.Input(key='valor')],
        [sg.Button('Voltar'), sg.Button('OK', key='add')]
    ]
    return sg.Window('Adicionar criança', layout=Adicionar, finalize=True)

def janela_mostrar():
    sg.theme('Reddit')
    Mostrar = [
        [sg.Output(size=(30,20))],
        [sg.Button('Voltar')]
    ]
    return sg.Window('Mostrar lista de crianças', layout=Mostrar, finalize=True)

def janela_retirar():
    sg.theme('Reddit')
    Retirar = [
        [sg.Text('Qual o numero da criança que será retirada?')],
        [sg.Input(key='reitirar_num')],
        [sg.Button('Voltar'), sg.Button('OK', key='pop')]
    ]
    return sg.Window('Remover criança', layout= Retirar, finalize=True)

#Janela
janela0, janela1, janela2, janela3, janela4 = janela_menu(), None, None, None, None

#ler ação
while True:
    janela, eventos, valores = sg.read_all_windows()
    if janela == janela0 and eventos == sg.WINDOW_CLOSED:
        break

    if janela == janela0 and eventos == '1- Fazer recibos':
        janela1 = janela_fazer()
        janela0.hide()
        document = Document()

        dia= (datetime.today().strftime('%d-%m-%Y'))
        styles = document.styles

        # stilo parágrafo
        p = styles.add_style('paragrafo', WD_STYLE_TYPE.PARAGRAPH)
        p.font.name = 'Arial'
        p.font.size = Pt(11)

        # stilo heading2
        h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
        h2.base_style = styles["Heading 2"]
        h2.font.name = "Arial"
        h2.font.size = Pt(11)
        h2.font.color.rgb = RGBColor(0, 0, 0)

        a = styles.add_style("assinatura", WD_STYLE_TYPE.PARAGRAPH)
        a.font.name = "Arial"
        a.font.size = Pt(10)
        a.font.color.rgb = RGBColor(0, 0, 0)
        a.font.bold = True

        r = styles.add_style("rodape", WD_STYLE_TYPE.PARAGRAPH)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0, 0, 0)

    if janela == janela1 and eventos == 'do':
        if janela == janela12 and eventos == 'Voltar':
            janela12.hide()
            janela0.un_hide()
        if janela == janela12 and eventos==sg.WINDOW_CLOSED:
            janela0.un_hide()
            janela12.close()

        recibo = 0
        clientes = []
        janela12 = janela_continuarfazer()
        janela1.hide()
        
        print('até aki foi 1')

        arquivo = open('info_recibo.txt', 'r', encoding='Utf-8', newline='')

        for linha in arquivo:
            linha = linha.strip(",")
            clientes.append(linha.split())
            
        arquivo.close()
        
        print('até aki foi 2')
        
        while True:
            if janela == janela12 and eventos == 'proximo':

                valor = (clientes[recibo+3])
                Total = valor*QtsDias
                printTotal= str(Total)
                QuaisDias = valores['quais']
                QtsDias = valores['dias']

                nomePai = clientes[recibo]
                cpfPai = clientes[recibo+1]
                nomeCrianca = clientes[recibo+2]
        
                print('até aki foi 3')
                
                document.add_paragraph('RECIBO                                                 VALOR: R$'+printTotal+',00', style="H2").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                document.add_paragraph('\n\nRecebi de '+nomePai+', CPF '+cpfPai+', o valor de '+printTotal+' reais, referente a '+QtsDias+' sessões de Terapia a '+nomeCrianca+', dias ('+QuaisDias+') de agosto de 2021.', style="paragrafo")
            
                document.add_paragraph('\nSão josé dos Campos,'+dia+'.', style="paragrafo").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                document.add_paragraph('\n___________________________\nJoâo da Silva \nTerapeuta \nCREFITO: xxxx-xx\nCPF: xxx.xxxx.xxx-xx', style="assinatura")
                document.add_paragraph('\n\nAlameda José Alves de Siqueira Filho, 52. Vila Betânia. CEP: 12245-492 São josé dos Campos-SP\ncel.(12)98812-9174 / e-mail. xxxxxxxxx@outlook.com\n\n\n ', style="rodape").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                

                document.save(nomeCrianca+'.docx')
                
                recibo += 4
            
            if janela == janela12 and eventos == 'finalizar':    
                sg.popup("Recibos prontos")
                break

    if janela == janela1 and eventos == 'Voltar':
        janela1.hide()
        janela0.un_hide()
    
    if janela == janela1 and eventos==sg.WINDOW_CLOSED:
        janela0.un_hide()
        janela1.close()

    if janela == janela0 and eventos == '2- Adicionar criança':
        janela2 = janela_adicionar()
        janela0.hide()

    if janela == janela2 and eventos == 'add':
        cpfpaivalido = ValidaCpf(valores['cpfpai'])
        #cpfcriancavalido = ValidaCpf(valores['cpfcrianca'])
        if cpfpaivalido.valida():
            tem = Ja_tem(valores['nomecrianca'])
            if tem == True:
                sg.popup("Essa criança já foi adicionada na lista")
            else:
                pai_nome = valores['nomepai']
                cpfpai = valores['cpfpai']
                criança_nome = valores['nomecrianca']
                #cpfcriança = valores['cpfcrianca']
                valorcobrado = valores['valor']
                with open('info_recibo.txt', 'a', encoding='Utf-8', newline='') as arquivo:
                    arquivo.writelines(f'{pai_nome} {cpfpai} {criança_nome} {valorcobrado}\n')
                    sg.popup('Pronto! ^-^')
        else:
            sg.popup("Informações incorretas!! Por favor verifique os CPF's")
    
    if janela == janela2 and eventos == 'Voltar':
        janela2.hide()
        janela0.un_hide()

    if janela == janela2 and eventos == sg.WINDOW_CLOSED:
        janela0.un_hide()
        janela2.close()

                   

    if janela == janela0 and eventos == '3- Mostrar lista de crianças':
        janela3 = janela_mostrar()
        janela0.hide()
        usuarios = []
        lista = 1
        with open('info_recibo.txt', 'r+', encoding='Utf-8', newline='') as arquivo:
            for linha in arquivo:
                linha = linha.strip(",")
                usuarios.append(linha.split())

            for usuario in usuarios:
                nome_crianca = usuario[2]
                print (lista, "- ",nome_crianca)
                lista+=1

    if janela == janela3 and eventos == 'Voltar':
        janela3.hide()
        janela0.un_hide()
    
    if janela == janela3 and eventos == sg.WINDOW_CLOSED:
        janela3.close()
        janela0.un_hide()


    if janela == janela0 and eventos == '4- Remover criança':
        janela4 = janela_retirar()
        janela0.hide()
    
    if janela == janela4 and eventos == 'pop':
        usuarios = []
        arquivo = open('info_recibo.txt', 'r', encoding='Utf-8', newline='')
        for linha in arquivo:
            linha = linha.strip(",")
            usuarios.append(linha.split())

        print(usuarios)
        
        arquivo = open('info_recibo.txt', 'w')   
        arquivo = open('info_recibo.txt', 'a', encoding='Utf-8', newline='')  

        retirar = int(valores['reitirar_num'])
        retirar -= 1
        print(retirar)
        usuarios.pop (retirar)

        for usuario in usuarios:
            nome_pai = usuario[0]
            cpf_pai = usuario[1]
            nome_crianca = usuario[2]
            #cpf_crianca = usuario[3]
            valor_cobrado = usuario[3]
            arquivo.write(f'{nome_pai} {cpf_pai} {nome_crianca} {valor_cobrado}\n')
        
        arquivo.close()
        sg.popup('Pronto! ^-^')  
          
    
    if janela == janela4 and eventos == 'Voltar':
        janela4.hide()
        janela0.un_hide()
    
    if janela == janela4 and eventos == sg.WINDOW_CLOSED:
        janela0.un_hide()
        janela4.close()


#Jeferson 093.785.730-02 Jonathan 142.825.950-35 200
#Alceu 048.265.290-07 Lucas 517.595.230-70 200
#Lucia 748.243.110-12 Luiz 554.095.090-99 200
#Marilucia 730.302.590-13 Amanda 672.189.600-14 200
#Fabio 089.916.730-66 Nicolas 399.565.680-31 200